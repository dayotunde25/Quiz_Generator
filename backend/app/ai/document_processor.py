"""
Document Processing Service
"""
import os
import re
import hashlib
from typing import Optional, Dict, Any
import magic
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import bleach


class DocumentProcessor:
    """Service for processing and extracting text from various document formats"""
    
    def __init__(self):
        self.max_file_size = 16 * 1024 * 1024  # 16MB
        self.allowed_extensions = {'txt', 'pdf', 'docx', 'doc', 'md', 'rtf'}
        self.allowed_mime_types = {
            'text/plain',
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
            'text/markdown',
            'application/rtf'
        }
    
    def validate_file(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Validate uploaded file"""
        result = {
            'is_valid': False,
            'error': None,
            'file_info': {}
        }
        
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                result['error'] = 'File not found'
                return result
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                result['error'] = f'File too large. Maximum size is {self.max_file_size // (1024*1024)}MB'
                return result
            
            if file_size == 0:
                result['error'] = 'File is empty'
                return result
            
            # Check file extension
            if '.' not in original_filename:
                result['error'] = 'File must have an extension'
                return result
            
            extension = original_filename.rsplit('.', 1)[1].lower()
            if extension not in self.allowed_extensions:
                result['error'] = f'File type not allowed. Allowed types: {", ".join(self.allowed_extensions)}'
                return result
            
            # Check MIME type
            try:
                mime_type = magic.from_file(file_path, mime=True)
                if mime_type not in self.allowed_mime_types:
                    result['error'] = f'Invalid file type detected: {mime_type}'
                    return result
            except Exception as e:
                print(f"MIME type detection failed: {e}")
                # Continue without MIME type check if magic fails
                mime_type = None
            
            # Calculate file hash for deduplication
            file_hash = self._calculate_file_hash(file_path)
            
            result['is_valid'] = True
            result['file_info'] = {
                'size': file_size,
                'extension': extension,
                'mime_type': mime_type,
                'hash': file_hash
            }
            
        except Exception as e:
            result['error'] = f'File validation error: {str(e)}'
        
        return result
    
    def extract_text(self, file_path: str, file_extension: str) -> Dict[str, Any]:
        """Extract text from document"""
        result = {
            'success': False,
            'text': '',
            'metadata': {},
            'error': None
        }
        
        try:
            if file_extension.lower() == 'txt':
                result = self._extract_from_txt(file_path)
            elif file_extension.lower() == 'pdf':
                result = self._extract_from_pdf(file_path)
            elif file_extension.lower() in ['docx', 'doc']:
                result = self._extract_from_docx(file_path)
            elif file_extension.lower() == 'md':
                result = self._extract_from_markdown(file_path)
            elif file_extension.lower() == 'rtf':
                result = self._extract_from_rtf(file_path)
            else:
                result['error'] = f'Unsupported file type: {file_extension}'
            
            # Clean and sanitize extracted text
            if result['success'] and result['text']:
                result['text'] = self._clean_text(result['text'])
                result['metadata']['word_count'] = len(result['text'].split())
                result['metadata']['character_count'] = len(result['text'])
            
        except Exception as e:
            result['error'] = f'Text extraction error: {str(e)}'
        
        return result
    
    def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            return {
                'success': True,
                'text': text,
                'metadata': {
                    'format': 'text',
                    'encoding': 'utf-8'
                }
            }
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': {},
                'error': str(e)
            }
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        text = ''
        metadata = {'format': 'pdf'}
        
        try:
            # Try with pdfplumber first (better for tables and layout)
            with pdfplumber.open(file_path) as pdf:
                metadata['page_count'] = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n\n'
            
            # If pdfplumber didn't extract much text, try PyMuPDF
            if len(text.strip()) < 100:
                doc = fitz.open(file_path)
                text = ''
                
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text += page.get_text() + '\n\n'
                
                doc.close()
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': metadata,
                'error': str(e)
            }
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ''
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + ' '
                    text += '\n'
            
            metadata = {
                'format': 'docx',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': {'format': 'docx'},
                'error': str(e)
            }
    
    def _extract_from_markdown(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            # Remove markdown formatting
            text = re.sub(r'#{1,6}\s+', '', text)  # Headers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Inline code
            text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)  # Code blocks
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links
            
            return {
                'success': True,
                'text': text,
                'metadata': {
                    'format': 'markdown'
                }
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': {'format': 'markdown'},
                'error': str(e)
            }
    
    def _extract_from_rtf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from RTF file"""
        try:
            # Basic RTF text extraction (simplified)
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # Remove RTF control words and formatting
            text = re.sub(r'\\[a-z]+\d*\s?', '', content)
            text = re.sub(r'[{}]', '', text)
            text = re.sub(r'\\\w+', '', text)
            
            return {
                'success': True,
                'text': text,
                'metadata': {
                    'format': 'rtf'
                }
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': {'format': 'rtf'},
                'error': str(e)
            }
    
    def _clean_text(self, text: str) -> str:
        """Clean and sanitize extracted text"""
        if not text:
            return ''
        
        # Remove HTML tags if any
        text = bleach.clean(text, tags=[], strip=True)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove excessive punctuation
        text = re.sub(r'[.]{3,}', '...', text)
        text = re.sub(r'[-]{3,}', '---', text)
        
        # Remove control characters
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        return text.strip()
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        
        return hash_sha256.hexdigest()
    
    def get_text_summary(self, text: str, max_length: int = 500) -> str:
        """Get a summary of the text"""
        if len(text) <= max_length:
            return text
        
        # Simple extractive summarization
        sentences = text.split('.')
        summary = ''
        
        for sentence in sentences:
            if len(summary + sentence) <= max_length:
                summary += sentence + '.'
            else:
                break
        
        return summary.strip()
